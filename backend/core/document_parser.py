"""
文档解析器 - 支持多种格式的合同文件解析
"""
import os
from pathlib import Path
from abc import ABC, abstractmethod
from typing import List, Optional, Dict, Any
import tempfile

from loguru import logger


class DocumentParser(ABC):
    """文档解析器抽象基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档，返回纯文本内容"""
        pass
    
    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文档元数据"""
        pass


class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析PDF文件"""
        try:
            import fitz  # pymupdf
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            doc.close()
            full_text = "\n\n".join(text_parts)
            logger.info(f"Parsed PDF: {file_path}, length: {len(full_text)} chars")
            return full_text
            
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            raise
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取PDF元数据"""
        try:
            import fitz
            
            doc = fitz.open(file_path)
            metadata = {
                "page_count": len(doc),
                "metadata": doc.metadata,
            }
            doc.close()
            return metadata
        except Exception as e:
            logger.error(f"Failed to get PDF metadata: {e}")
            return {}


class WordParser(DocumentParser):
    """Word文档解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析Word文件"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"Parsed Word: {file_path}, length: {len(full_text)} chars")
            return full_text
            
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Failed to parse Word document: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取Word文档元数据"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            return {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        except Exception as e:
            logger.error(f"Failed to get Word metadata: {e}")
            return {}


class TextParser(DocumentParser):
    """纯文本解析器"""
    
    def parse(self, file_path: str) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"Parsed text: {file_path}, length: {len(text)} chars")
            return text
        except Exception as e:
            logger.error(f"Failed to read text file: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文本文件元数据"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "lines": sum(1 for _ in open(file_path, 'r')),
            }
        except Exception as e:
            logger.error(f"Failed to get text metadata: {e}")
            return {}


class DocumentParserFactory:
    """文档解析器工厂"""
    
    _parsers = {
        '.pdf': PDFParser(),
        '.docx': WordParser(),
        '.doc': WordParser(),
        '.txt': TextParser(),
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """根据文件扩展名获取解析器"""
        ext = Path(file_path).suffix.lower()
        
        if ext in cls._parsers:
            return cls._parsers[ext]
        
        raise ValueError(f"Unsupported file format: {ext}. Supported: {list(cls._parsers.keys())}")
    
    @classmethod
    def parse_document(cls, file_path: str) -> str:
        """解析文档"""
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)
    
    @classmethod
    def get_document_metadata(cls, file_path: str) -> Dict[str, Any]:
        """获取文档元数据"""
        parser = cls.get_parser(file_path)
        return parser.get_metadata(file_path)
    
    @classmethod
    def supported_formats(cls) -> List[str]:
        """支持的格式"""
        return list(cls._parsers.keys())


# 便捷函数
def parse_document(file_path: str) -> str:
    """解析文档"""
    return DocumentParserFactory.parse_document(file_path)


def get_document_metadata(file_path: str) -> Dict[str, Any]:
    """获取文档元数据"""
    return DocumentParserFactory.get_document_metadata(file_path)
